"""
NEA Code Documentation Generator — Flutter/Dart Edition
========================================================
Generates a comprehensive Microsoft Word (.docx) document that fully documents
the Flutter/Dart mobile application for AQA A-Level Computer Science NEA submission.

This script:
- Parses Dart source files (classes, methods, constructors, mixins)
- Preserves ALL original code comments verbatim
- Generates vertical documentation tables (two columns: label → value)
- Adds explanatory TEXTBOX ANNOTATIONS in Word (JSON-driven)
- Loads ALL explanatory text from external metadata files
- Applies DIFFERENT documentation styles depending on file type:
    * Services/ML/Data Structures: Full NEA tables with AQA skill mapping
    * Views/UI: Simplified widget tables
    * Models/Constants: Brief file-level tables

Adapted from the Flask NEA Documentation Generator for a Flutter/Dart codebase.

Author: NEA Submission Documentation Generator
Date: February 2026
"""

import os
import sys
import json
import re
import argparse
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class MissingMetadataError(Exception):
    """Raised when required metadata is missing from JSON files."""
    pass


class DartParser:
    """
    Lightweight Dart source code parser.
    
    Since Dart does not have a Python-style AST module, this parser uses
    regular expressions and brace-counting to extract classes, methods,
    constructors, and top-level functions from Dart source files.
    
    This is NOT a full Dart parser — it handles the patterns found in this
    specific codebase reliably. For AQA NEA purposes, this demonstrates
    string processing, regex pattern matching, and algorithmic parsing.
    """

    @staticmethod
    def extract_imports(source_code: str) -> List[Dict]:
        """
        Extract all import statements from Dart source code.
        
        Args:
            source_code: The full Dart source file contents.
            
        Returns:
            A list of dicts with 'line', 'lineno', and 'package' keys.
        """
        imports = []
        for i, line in enumerate(source_code.split('\n'), 1):
            stripped = line.strip()
            if stripped.startswith('import ') or stripped.startswith("import '"):
                # Extract the package path from the import
                match = re.search(r"import\s+'([^']+)'", stripped)
                package = match.group(1) if match else stripped
                imports.append({
                    'line': stripped,
                    'lineno': i,
                    'package': package
                })
        return imports

    @staticmethod
    def extract_classes(source_code: str) -> List[Dict]:
        """
        Extract class definitions with their methods, constructors, and fields.
        
        Uses brace-counting to find class boundaries, then parses the class
        body for method signatures. This approach handles nested braces correctly.
        
        Args:
            source_code: The full Dart source file contents.
            
        Returns:
            A list of class info dicts containing name, methods, fields, etc.
        """
        classes = []
        lines = source_code.split('\n')
        
        # Pattern to match class declarations (including abstract, with mixins, extends)
        class_pattern = re.compile(
            r'^(abstract\s+)?class\s+(\w+)'
            r'(?:\s+extends\s+([\w<>,\s]+?))?'
            r'(?:\s+with\s+([\w<>,\s]+?))?'
            r'(?:\s+implements\s+([\w<>,\s]+?))?'
            r'\s*\{'
        )
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            match = class_pattern.match(line)
            if match:
                is_abstract = match.group(1) is not None
                class_name = match.group(2)
                extends_class = match.group(3).strip() if match.group(3) else None
                with_mixins = match.group(4).strip() if match.group(4) else None
                implements = match.group(5).strip() if match.group(5) else None
                
                # Find the class body using brace counting
                start_line = i
                brace_count = 0
                class_body_lines = []
                j = i
                found_start = False
                
                while j < len(lines):
                    for ch in lines[j]:
                        if ch == '{':
                            brace_count += 1
                            found_start = True
                        elif ch == '}':
                            brace_count -= 1
                    
                    class_body_lines.append(lines[j])
                    
                    if found_start and brace_count == 0:
                        break
                    j += 1
                
                end_line = j
                class_body = '\n'.join(class_body_lines)
                
                # Extract methods from the class body
                methods = DartParser._extract_methods_from_body(
                    class_body_lines[1:], start_line + 2, class_name
                )
                
                # Extract class-level fields
                fields = DartParser._extract_fields(class_body_lines[1:])
                
                # Get class docstring (comments above class declaration)
                docstring = DartParser._get_docstring(lines, start_line)
                
                classes.append({
                    'name': class_name,
                    'is_abstract': is_abstract,
                    'extends': extends_class,
                    'with_mixins': with_mixins,
                    'implements': implements,
                    'lineno': start_line + 1,
                    'end_lineno': end_line + 1,
                    'docstring': docstring,
                    'methods': methods,
                    'fields': fields,
                    'source': class_body
                })
                
                i = end_line + 1
            else:
                i += 1
        
        return classes

    @staticmethod
    def _extract_methods_from_body(body_lines: List[str], start_offset: int, class_name: str) -> List[Dict]:
        """
        Extract method and constructor definitions from a class body.
        
        Args:
            body_lines: Lines of code inside the class braces.
            start_offset: Line number offset for accurate line reporting.
            class_name: Name of the containing class.
            
        Returns:
            List of method info dicts.
        """
        methods = []
        
        # Patterns for different method types
        # Regular methods: returnType methodName(params) { or async {
        method_pattern = re.compile(
            r'^\s+(?:static\s+)?(?:Future<[^>]+>|Stream<[^>]+>|void|int|double|bool|String|List<[^>]+>|Map<[^>]+>|Set<[^>]+>|[\w<>,\s]+?)\s+'
            r'(\w+)\s*\(([^)]*)\)\s*(?:async\s*)?\{'
        )
        
        # Constructors: ClassName(params) or ClassName.named(params)
        constructor_pattern = re.compile(
            rf'^\s+{re.escape(class_name)}(?:\.(\w+))?\s*\(([^)]*)\)'
        )
        
        # Factory constructors
        factory_pattern = re.compile(
            rf'^\s+factory\s+{re.escape(class_name)}(?:\.(\w+))?\s*\(([^)]*)\)'
        )
        
        # Getter pattern
        getter_pattern = re.compile(
            r'^\s+(?:[\w<>,\s]+?)\s+get\s+(\w+)\s*(?:\{|=>)'
        )
        
        # Override annotation
        override_pattern = re.compile(r'^\s+@override\s*$')
        
        i = 0
        while i < len(body_lines):
            line = body_lines[i].strip()
            
            # Skip empty lines and comments
            if not line or line.startswith('//') or line.startswith('/*') or line.startswith('*'):
                i += 1
                continue
            
            full_line = body_lines[i]
            
            # Check for factory constructor
            factory_match = factory_pattern.match(full_line)
            if factory_match:
                named = factory_match.group(1) or ''
                params_str = factory_match.group(2)
                method_name = f'factory {class_name}' + (f'.{named}' if named else '')
                
                # Get docstring
                docstring = DartParser._get_docstring(body_lines, i)
                
                # Find method end
                end_i = DartParser._find_method_end(body_lines, i)
                
                methods.append({
                    'name': method_name,
                    'type': 'factory_constructor',
                    'params_raw': params_str.strip(),
                    'lineno': start_offset + i,
                    'end_lineno': start_offset + end_i,
                    'docstring': docstring,
                    'is_override': False,
                    'source': '\n'.join(body_lines[i:end_i+1])
                })
                
                i = end_i + 1
                continue
            
            # Check for constructor
            constructor_match = constructor_pattern.match(full_line)
            if constructor_match and not line.startswith('return') and not line.startswith('if'):
                named = constructor_match.group(1) or ''
                params_str = constructor_match.group(2)
                method_name = class_name + (f'.{named}' if named else '')
                
                docstring = DartParser._get_docstring(body_lines, i)
                end_i = DartParser._find_method_end(body_lines, i)
                
                methods.append({
                    'name': method_name,
                    'type': 'constructor',
                    'params_raw': params_str.strip(),
                    'lineno': start_offset + i,
                    'end_lineno': start_offset + end_i,
                    'docstring': docstring,
                    'is_override': False,
                    'source': '\n'.join(body_lines[i:end_i+1])
                })
                
                i = end_i + 1
                continue
            
            # Check for getter
            getter_match = getter_pattern.match(full_line)
            if getter_match:
                method_name = f'get {getter_match.group(1)}'
                docstring = DartParser._get_docstring(body_lines, i)
                end_i = DartParser._find_method_end(body_lines, i)
                
                methods.append({
                    'name': method_name,
                    'type': 'getter',
                    'params_raw': '',
                    'lineno': start_offset + i,
                    'end_lineno': start_offset + end_i,
                    'docstring': docstring,
                    'is_override': False,
                    'source': '\n'.join(body_lines[i:end_i+1])
                })
                
                i = end_i + 1
                continue
            
            # Check for regular method (including void, Future, static, etc.)
            # More flexible pattern
            flexible_method = re.match(
                r'^\s+(?:@\w+\s+)?(?:static\s+)?(?:[\w<>\?,\s]+?\s+)?(\w+)\s*\(([^)]*)\)\s*(?:async\s*)?(?:\{|=>)',
                full_line
            )
            if flexible_method:
                method_name = flexible_method.group(1)
                params_str = flexible_method.group(2)
                
                # Skip if it's a control structure
                if method_name in ('if', 'for', 'while', 'switch', 'catch', 'return', 'setState', 'then', 'where', 'map', 'forEach', 'when', 'showDialog', 'Navigator', 'ScaffoldMessenger'):
                    i += 1
                    continue
                
                # Check for @override on previous line
                is_override = False
                if i > 0 and body_lines[i-1].strip() == '@override':
                    is_override = True
                
                docstring = DartParser._get_docstring(body_lines, i)
                end_i = DartParser._find_method_end(body_lines, i)
                
                # Determine method type
                is_static = 'static ' in full_line
                is_async = 'async' in full_line
                
                methods.append({
                    'name': method_name,
                    'type': 'static_method' if is_static else 'method',
                    'params_raw': params_str.strip(),
                    'lineno': start_offset + i,
                    'end_lineno': start_offset + end_i,
                    'docstring': docstring,
                    'is_override': is_override,
                    'is_async': is_async,
                    'source': '\n'.join(body_lines[i:end_i+1])
                })
                
                i = end_i + 1
                continue
            
            i += 1
        
        return methods

    @staticmethod
    def _find_method_end(lines: List[str], start: int) -> int:
        """Find the end of a method/function body using brace counting."""
        brace_count = 0
        found_start = False
        
        # Check for arrow functions (=>)
        first_line = lines[start] if start < len(lines) else ''
        if '=>' in first_line and '{' not in first_line.split('=>')[0]:
            # Arrow function — find the semicolon
            j = start
            while j < len(lines):
                if ';' in lines[j]:
                    return j
                j += 1
            return min(start + 1, len(lines) - 1)
        
        j = start
        while j < len(lines):
            for ch in lines[j]:
                if ch == '{':
                    brace_count += 1
                    found_start = True
                elif ch == '}':
                    brace_count -= 1
            
            if found_start and brace_count <= 0:
                return j
            j += 1
        
        return min(start + 5, len(lines) - 1)

    @staticmethod
    def _extract_fields(body_lines: List[str]) -> List[str]:
        """Extract class-level field declarations."""
        fields = []
        field_pattern = re.compile(
            r'^\s+(?:final\s+|late\s+|static\s+|const\s+)*'
            r'(?:[\w<>\?,\s]+?)\s+'
            r'(\w+)\s*[;=]'
        )
        
        for line in body_lines:
            stripped = line.strip()
            # Skip methods, comments, annotations
            if stripped.startswith('//') or stripped.startswith('@') or stripped.startswith('*'):
                continue
            if '(' in stripped and ')' in stripped:
                continue
                
            match = field_pattern.match(line)
            if match:
                field_name = match.group(1)
                if field_name not in ('if', 'for', 'while', 'return', 'class', 'void'):
                    fields.append(field_name)
        
        return fields

    @staticmethod
    def _get_docstring(lines: List[str], target_line: int) -> str:
        """
        Get the documentation comment (/// or /** */) above a declaration.
        
        Args:
            lines: All source lines.
            target_line: The line index of the declaration.
            
        Returns:
            Concatenated doc comment text, or empty string.
        """
        doc_lines = []
        i = target_line - 1
        
        while i >= 0:
            stripped = lines[i].strip()
            if stripped.startswith('///'):
                doc_lines.insert(0, stripped[3:].strip())
            elif stripped.startswith('*') or stripped.startswith('/**'):
                doc_lines.insert(0, stripped.lstrip('/* ').rstrip('*/').strip())
            elif stripped == '' or stripped.startswith('@'):
                i -= 1
                continue
            else:
                break
            i -= 1
        
        return ' '.join(doc_lines).strip()

    @staticmethod
    def extract_top_level_functions(source_code: str, class_ranges: List[Tuple[int, int]]) -> List[Dict]:
        """
        Extract top-level (non-class) functions.
        
        Args:
            source_code: Full source code.
            class_ranges: List of (start, end) line ranges occupied by classes.
            
        Returns:
            List of function info dicts.
        """
        functions = []
        lines = source_code.split('\n')
        
        func_pattern = re.compile(
            r'^(?:Future<[^>]+>|Stream<[^>]+>|void|int|double|bool|String|List<[^>]+>|Map<[^>]+>|[\w<>,\s]+?)\s+'
            r'(\w+)\s*\(([^)]*)\)\s*(?:async\s*)?\{'
        )
        
        i = 0
        while i < len(lines):
            # Skip if inside a class
            in_class = any(start <= i + 1 <= end for start, end in class_ranges)
            if in_class:
                i += 1
                continue
            
            line = lines[i]
            if line and not line[0].isspace():  # Top-level (no indentation)
                match = func_pattern.match(line)
                if match:
                    func_name = match.group(1)
                    params_str = match.group(2)
                    
                    if func_name in ('if', 'for', 'while', 'class', 'enum', 'import', 'export', 'mixin'):
                        i += 1
                        continue
                    
                    docstring = DartParser._get_docstring(lines, i)
                    end_i = DartParser._find_method_end(lines, i)
                    
                    functions.append({
                        'name': func_name,
                        'params_raw': params_str.strip(),
                        'lineno': i + 1,
                        'end_lineno': end_i + 1,
                        'docstring': docstring,
                        'source': '\n'.join(lines[i:end_i+1])
                    })
                    
                    i = end_i + 1
                    continue
            
            i += 1
        
        return functions

    @staticmethod
    def extract_enums(source_code: str) -> List[Dict]:
        """Extract enum definitions from Dart source code."""
        enums = []
        lines = source_code.split('\n')
        
        enum_pattern = re.compile(r'^enum\s+(\w+)\s*\{')
        
        i = 0
        while i < len(lines):
            match = enum_pattern.match(lines[i].strip())
            if match:
                enum_name = match.group(1)
                end_i = DartParser._find_method_end(lines, i)
                docstring = DartParser._get_docstring(lines, i)
                
                enums.append({
                    'name': enum_name,
                    'lineno': i + 1,
                    'end_lineno': end_i + 1,
                    'docstring': docstring,
                    'source': '\n'.join(lines[i:end_i+1])
                })
                
                i = end_i + 1
            else:
                i += 1
        
        return enums

    @staticmethod
    def extract_extensions(source_code: str) -> List[Dict]:
        """Extract extension definitions from Dart source code."""
        extensions = []
        lines = source_code.split('\n')
        
        ext_pattern = re.compile(r'^extension\s+(\w+)(?:<[^>]+>)?\s+on\s+(.+?)\s*\{')
        
        i = 0
        while i < len(lines):
            match = ext_pattern.match(lines[i].strip())
            if match:
                ext_name = match.group(1)
                on_type = match.group(2).strip()
                end_i = DartParser._find_method_end(lines, i)
                docstring = DartParser._get_docstring(lines, i)
                
                extensions.append({
                    'name': ext_name,
                    'on_type': on_type,
                    'lineno': i + 1,
                    'end_lineno': end_i + 1,
                    'docstring': docstring,
                    'source': '\n'.join(lines[i:end_i+1])
                })
                
                i = end_i + 1
            else:
                i += 1
        
        return extensions


# =============================================================================
# MAIN DOCUMENTATION GENERATOR
# =============================================================================

class NEACodeDocumentationGenerator:
    """
    Generates comprehensive NEA code documentation for AQA A-Level CS.
    
    Adapted for Flutter/Dart codebase. This generator follows strict rules:
    - All explanations come from metadata files (single source of truth)
    - Code is preserved verbatim with all comments
    - Different file types get different documentation styles
    - Vertical tables only (no horizontal tables)
    - British spelling throughout
    """

    # File type classifications for Dart/Flutter
    SERVICE_FILES = [
        'call_cycle_service.dart', 'excel_import_service.dart', 'team_service.dart',
        'firebase_auth_provider.dart', 'auth_provider.dart', 'auth_service.dart',
        'firebase_cloud_storage.dart', 'notes_service.dart',
    ]
    ML_FILES = [
        'kaplan_meier.dart', 'linear_regression.dart',
        'naive_bayes_classifier.dart', 'wilson_score.dart',
    ]
    DATA_STRUCTURE_FILES = [
        'contact_priority_queue.dart', 'trie.dart', 'nlp_service.dart',
    ]
    MODEL_FILES = [
        'app_user.dart', 'team.dart', 'auth_user.dart', 'auth_exceptions.dart',
        'cloud_note.dart', 'cloud_list.dart',
        'cloud_storage_constants.dart',
    ]
    VIEW_FILES_PATTERNS = [
        'views/', 'new-login-pages/',
    ]
    UTILITY_FILES_PATTERNS = [
        'utilities/', 'helpers/', 'extensions/', 'enums/', 'theme/',
        'constants/',
    ]
    
    # Files to skip entirely
    SKIP_FILES = [
        'firebase_options.dart',  # Auto-generated
    ]
    
    # Files that are commented out / empty
    COMMENTED_OUT_FILES = [
        'notes_service.dart', 'call_service.dart', 'call2_service.dart',
        'crud_exceptions.dart', 'dialer_contacts_view.dart',
    ]

    def __init__(self, project_root: str, output_file: str = "NEA_Code_Documentation.docx",
                 append_mode: bool = False):
        """
        Initialise the documentation generator.
        
        Args:
            project_root: Path to the project root directory.
            output_file: Name of the output Word document.
            append_mode: If True, append to existing document.
        """
        self.project_root = Path(project_root)
        self.output_file = output_file
        self.append_mode = append_mode
        
        # Determine workspace root (parent of lib if project_root ends with lib)
        if self.project_root.name == 'lib':
            self.workspace_root = self.project_root.parent
        else:
            self.workspace_root = self.project_root
        
        # Load existing document or create new one
        output_path = self.project_root / output_file
        if append_mode and output_path.exists():
            self.doc = Document(str(output_path))
            self.doc.add_page_break()
            print(f"Appending to existing document: {output_path}")
        else:
            self.doc = Document()
        
        # Load metadata from files (relative to workspace root, not project_root)
        self.class_explanations = self._load_all_metadata()
        self.file_metadata = self._load_file_level_metadata()
        self.system_context = self._load_json('lib/ExplanationFiles/system_context.json')
        self.test_data = self._load_json('lib/ExplanationFiles/test_data.json')
        self.project_objectives = self._load_project_objectives()
        
        # Track missing metadata for reporting
        self.missing_metadata = []
        
        # Setup document styles
        self._setup_styles()

    # =========================================================================
    # METADATA LOADING
    # =========================================================================

    def _load_json(self, relative_path: str) -> Dict:
        """Load a JSON file safely."""
        filepath = self.workspace_root / relative_path
        if filepath.exists():
            with open(filepath, 'r', encoding='utf-8') as f:
                return json.load(f)
        print(f"Warning: Could not load {filepath}")
        return {}

    def _load_project_objectives(self) -> List[Dict]:
        """Load project objectives from the ExtraFiles folder."""
        obj_path = self.workspace_root / 'ExtraFiles' / 'projectObjectives.txt'
        objectives = []
        if obj_path.exists():
            with open(obj_path, 'r', encoding='utf-8') as f:
                content = f.read()
            # Parse objectives — they follow pattern "X.Y description"
            current_category = ""
            for line in content.split('\n'):
                line = line.strip()
                if line.startswith('Category'):
                    current_category = line
                elif line and line[0].isdigit() and '\t' in line:
                    parts = line.split('\t', 1)
                    obj_id = parts[0].strip()
                    desc = parts[1].strip() if len(parts) > 1 else ''
                    objectives.append({
                        'id': obj_id,
                        'description': desc,
                        'category': current_category
                    })
        return objectives

    def _load_all_metadata(self) -> Dict[str, Dict]:
        """
        Load and merge metadata from all sources.
        
        Primary source: lib/ExplanationFiles/class_explanations_extended.txt
        Format for each entry:
            [ClassName.methodName]
            PURPOSE: What this method does and why
            PARAMETERS:
            - paramName (type): explanation
            MARKSCHEME: Group X | Skill Name | How Applied
            OBJECTIVE: X.Y - Description
        
        Returns:
            Comprehensive dictionary keyed by 'ClassName.methodName'.
        """
        metadata = {}
        
        ext_paths = [
            self.workspace_root / 'lib' / 'ExplanationFiles' / 'class_explanations_extended.txt',
        ]
        
        for ext_path in ext_paths:
            if ext_path.exists():
                current_id = None
                current_data = {}
                current_field = None
                current_params = []
                pending_lines = []
                
                with open(ext_path, 'r', encoding='utf-8') as f:
                    lines = f.readlines()
                
                for i, line in enumerate(lines):
                    line_stripped = line.strip()
                    
                    # Skip empty lines and separators
                    if not line_stripped or line_stripped.startswith('=') or line_stripped.startswith('━'):
                        continue
                    
                    # Skip section headers and file-level metadata lines
                    if line_stripped.startswith('Author:') or line_stripped.startswith('Date:') or line_stripped.startswith('This document'):
                        continue
                    if line_stripped.startswith('FILE PURPOSE:') or line_stripped.startswith('WHY THIS FILE EXISTS:') or line_stripped.startswith('ROLE IN PROJECT:') or line_stripped.startswith('IMPORTS EXPLAINED:'):
                        current_field = None  # Stop appending to previous entry
                        continue
                    
                    # Match [ClassName.methodName] or [functionName] or [ClassName.factory ClassName]
                    match = re.match(r'^\[([A-Za-z0-9_ ]+(?:\.[A-Za-z0-9_ ]*)?)\]\s*(?:\(([^)]+)\))?$', line_stripped)
                    if match:
                        # Save previous entry
                        if current_id:
                            if pending_lines and not current_data.get('purpose'):
                                current_data['purpose'] = ' '.join(pending_lines).strip()
                            if current_params:
                                current_data['parameters_detailed'] = current_params
                            metadata[current_id] = current_data
                        
                        current_id = match.group(1)
                        location_hint = match.group(2)
                        current_data = {
                            'purpose': '',
                            'parameters_detailed': [],
                            'group': '',
                            'skill_wording': '',
                            'how_applied': '',
                            'objective': '',
                            'location': location_hint or '',
                            'source': 'extended'
                        }
                        current_params = []
                        current_field = None
                        pending_lines = []
                        continue
                    
                    if not current_id:
                        continue
                    
                    # Parse structured fields
                    if line_stripped.startswith('PURPOSE:'):
                        current_field = 'purpose'
                        purpose_text = line_stripped[8:].strip()
                        if purpose_text:
                            current_data['purpose'] = purpose_text
                    elif line_stripped.startswith('PARAMETERS:'):
                        current_field = 'parameters'
                        current_params = []
                    elif line_stripped.startswith('MARKSCHEME:'):
                        current_field = 'markscheme'
                        markscheme_text = line_stripped[11:].strip()
                        parts = [p.strip() for p in markscheme_text.split('|')]
                        if len(parts) >= 1:
                            current_data['group'] = parts[0]
                        if len(parts) >= 2:
                            current_data['skill_wording'] = parts[1]
                        if len(parts) >= 3:
                            current_data['how_applied'] = parts[2]
                    elif line_stripped.startswith('OBJECTIVE:'):
                        current_field = 'objective'
                        current_data['objective'] = line_stripped[10:].strip()
                    elif line_stripped.startswith('  - ') or line_stripped.startswith('- '):
                        if current_field == 'parameters':
                            param_line = line_stripped.lstrip('- ').strip()
                            param_match = re.match(r'([A-Za-z0-9_]+)\s*(?:\(([^)]+)\))?:\s*(.+)', param_line)
                            if param_match:
                                current_params.append({
                                    'name': param_match.group(1),
                                    'type': param_match.group(2) or 'dynamic',
                                    'explanation': param_match.group(3)
                                })
                            else:
                                current_params.append({
                                    'name': param_line.split(':')[0].strip(),
                                    'type': '',
                                    'explanation': param_line
                                })
                    elif current_field == 'purpose':
                        current_data['purpose'] += ' ' + line_stripped
                    elif current_field == 'objective':
                        current_data['objective'] += ' ' + line_stripped
                    else:
                        pending_lines.append(line_stripped)
                
                # Save last entry
                if current_id:
                    if pending_lines and not current_data.get('purpose'):
                        current_data['purpose'] = ' '.join(pending_lines).strip()
                    if current_params:
                        current_data['parameters_detailed'] = current_params
                    metadata[current_id] = current_data
                
                print(f"Loaded {len(metadata)} metadata entries from {ext_path}")
        
        return metadata

    def _load_file_level_metadata(self) -> Dict[str, Dict]:
        """
        Load file-level metadata (FILE PURPOSE, IMPORTS EXPLAINED, etc.).
        
        Returns:
            Dictionary keyed by normalised filename.
        """
        file_metadata = {}
        
        ext_path = self.workspace_root / 'lib' / 'ExplanationFiles' / 'class_explanations_extended.txt'
        
        if ext_path.exists():
            with open(ext_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Find all file sections using regex
            section_pattern = r'===\s*([A-Za-z0-9_/]+\.dart)\s*(?:\([^)]+\))?\s*==='
            sections = re.split(section_pattern, content, flags=re.IGNORECASE)
            
            for i in range(1, len(sections), 2):
                if i + 1 < len(sections):
                    filename = sections[i].lower()
                    section_content = sections[i + 1]
                    
                    file_data = {
                        'file_purpose': '',
                        'why_exists': '',
                        'role_in_project': '',
                        'imports_explained': {}
                    }
                    
                    # Extract FILE PURPOSE
                    purpose_match = re.search(
                        r'FILE PURPOSE:\s*(.+?)(?=\n\n|\nWHY|\nROLE|\nIMPORTS|\n===|\Z)',
                        section_content, re.DOTALL
                    )
                    if purpose_match:
                        file_data['file_purpose'] = ' '.join(purpose_match.group(1).split())
                    
                    # Extract WHY THIS FILE EXISTS
                    why_match = re.search(
                        r'WHY THIS FILE EXISTS:\s*(.+?)(?=\n\n|\nROLE|\nIMPORTS|\n\[|\n===|\Z)',
                        section_content, re.DOTALL
                    )
                    if why_match:
                        file_data['why_exists'] = ' '.join(why_match.group(1).split())
                    
                    # Extract ROLE IN PROJECT
                    role_match = re.search(
                        r'ROLE IN PROJECT:\s*(.+?)(?=\n\n|\n\[|\nIMPORTS|\n===|\Z)',
                        section_content, re.DOTALL
                    )
                    if role_match:
                        file_data['role_in_project'] = ' '.join(role_match.group(1).split())
                    
                    # Extract IMPORTS EXPLAINED
                    imports_match = re.search(
                        r'IMPORTS EXPLAINED:\s*\n(.+?)(?=\n\n===|\n===|\Z)',
                        section_content, re.DOTALL
                    )
                    if imports_match:
                        for import_line in imports_match.group(1).split('\n'):
                            import_line = import_line.strip().lstrip('- ')
                            import_m = re.match(r'([A-Za-z0-9_./]+(?:\s*\([^)]+\))?)\s*:\s*(.+)', import_line)
                            if import_m:
                                file_data['imports_explained'][import_m.group(1).strip()] = import_m.group(2).strip()
                    
                    file_metadata[filename] = file_data
            
            print(f"Loaded file-level metadata for {len(file_metadata)} files")
        
        return file_metadata

    # =========================================================================
    # DOCUMENT STYLES
    # =========================================================================

    def _setup_styles(self):
        """Configure professional document styles for NEA submission."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        styles = self.doc.styles
        
        try:
            code_char = styles.add_style('CodeChar', WD_STYLE_TYPE.CHARACTER)
            code_char.font.name = 'Cascadia Code'
            code_char.font.size = Pt(9)
        except ValueError:
            pass
        
        try:
            code_para = styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
            code_para.font.name = 'Cascadia Code'
            code_para.font.size = Pt(9)
            code_para.paragraph_format.left_indent = Cm(0.3)
            code_para.paragraph_format.space_before = Pt(3)
            code_para.paragraph_format.space_after = Pt(3)
            code_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        except ValueError:
            pass
        
        try:
            table_label = styles.add_style('TableLabel', WD_STYLE_TYPE.PARAGRAPH)
            table_label.font.name = 'Calibri'
            table_label.font.size = Pt(10)
            table_label.font.bold = True
        except ValueError:
            pass

    # =========================================================================
    # TABLE & FORMATTING HELPERS
    # =========================================================================

    def _set_cell_background(self, cell, color_hex: str):
        """Set table cell background colour using XML."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    def _set_cell_width(self, cell, width_cm: float):
        """Set the width of a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(width_cm * 567)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    def _get_group_colour(self, group: str) -> str:
        """Get background colour for AQA markscheme group."""
        group_upper = group.upper() if group else ''
        if 'A' in group_upper:
            return "C6EFCE"  # Light green — Group A (complex skills)
        elif 'B' in group_upper:
            return "FFEB9C"  # Light amber — Group B (intermediate skills)
        elif 'C' in group_upper:
            return "BDD7EE"  # Light blue — Group C (foundation skills)
        return "FFFFFF"

    def _add_heading(self, text: str, level: int):
        """Add a heading with consistent formatting."""
        return self.doc.add_heading(text, level=level)

    def _add_paragraph(self, text: str, bold: bool = False, italic: bool = False):
        """Add a paragraph with optional formatting."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        return p

    def _create_vertical_table(self, rows: List[Tuple[str, str]],
                                label_width_cm: float = 4.0,
                                value_width_cm: float = 12.0,
                                highlight_rows: Dict[str, str] = None) -> Any:
        """Create a vertical two-column table (label | value)."""
        if highlight_rows is None:
            highlight_rows = {}
        
        table = self.doc.add_table(rows=len(rows), cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        for i, (label, value) in enumerate(rows):
            label_cell = table.rows[i].cells[0]
            value_cell = table.rows[i].cells[1]
            
            self._set_cell_background(label_cell, "E8E8E8")
            self._set_cell_width(label_cell, label_width_cm)
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.bold = True
            label_run.font.size = Pt(10)
            
            if label in highlight_rows:
                self._set_cell_background(value_cell, highlight_rows[label])
            
            self._set_cell_width(value_cell, value_width_cm)
            
            # For long text values, split into paragraphs at sentence boundaries
            value_str = str(value)
            segments = self._split_into_paragraphs(value_str, label)
            
            for seg_idx, segment in enumerate(segments):
                if seg_idx == 0:
                    value_para = value_cell.paragraphs[0]
                else:
                    value_para = value_cell.add_paragraph()
                    value_para.paragraph_format.space_before = Pt(4)
                value_run = value_para.add_run(segment)
                value_run.font.size = Pt(10)
                
                if label in ["Markscheme Group", "Markscheme Skill", "How Applied"]:
                    value_run.bold = True
        
        return table

    def _split_into_paragraphs(self, text: str, label: str = "") -> List[str]:
        """Split a long text value into logical paragraphs for cleaner table cells.
        
        For short text (<200 chars), returns as-is.
        For longer text, splits at sentence boundaries ('. ') so each idea
        gets its own paragraph in the Word cell.
        Explicit newlines are always honoured.
        """
        # Always honour explicit newlines first
        if '\n' in text:
            return [seg.strip() for seg in text.split('\n') if seg.strip()]
        
        # Short text — no splitting needed
        if len(text) < 200:
            return [text]
        
        # Don't split certain fields that should stay compact
        compact_labels = {"Identifier", "File Path", "Source File", "File Type",
                          "Lines of Code", "Documentation Style", "Class Name",
                          "Widget Name", "Abstract", "Methods Count", "Returns",
                          "Markscheme Group", "Markscheme Skill", "Type"}
        if label in compact_labels:
            return [text]
        
        # Split at sentence boundaries (period followed by space and uppercase)
        # Also split at em-dashes used as section separators
        segments = []
        current = []
        sentences = re.split(r'(?<=\.)\s+(?=[A-Z])', text)
        
        for sentence in sentences:
            current.append(sentence)
            joined = ' '.join(current)
            # Start a new paragraph roughly every 1-3 sentences
            if len(joined) >= 150:
                segments.append(joined)
                current = []
        
        if current:
            segments.append(' '.join(current))
        
        # If splitting produced only 1 segment, return as-is
        if len(segments) <= 1:
            return [text]
        
        return segments

    def _add_code_to_cell(self, cell, code: str):
        """Add code-formatted text to a table cell."""
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(code)
        run.font.name = 'Cascadia Code'
        run.font.size = Pt(7.5)

    def _add_textbox_annotation(self, text: str, reference: str = ""):
        """Add a textbox annotation (simulated with bordered table cell)."""
        table = self.doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        
        cell = table.rows[0].cells[0]
        self._set_cell_background(cell, "FFF3CD")
        
        if reference:
            ref_para = cell.paragraphs[0]
            ref_run = ref_para.add_run(f"Annotation: {reference}")
            ref_run.bold = True
            ref_run.font.size = Pt(9)
            ref_run.font.color.rgb = RGBColor(0x85, 0x64, 0x04)
            
            text_para = cell.add_paragraph()
            text_run = text_para.add_run(text)
            text_run.font.size = Pt(9)
            text_run.italic = True
        else:
            para = cell.paragraphs[0]
            run = para.add_run(text)
            run.font.size = Pt(9)
            run.italic = True
        
        self.doc.add_paragraph()

    # =========================================================================
    # METADATA LOOKUP
    # =========================================================================

    def _get_method_explanation(self, class_name: str, method_name: str,
                                 item_type: str = "method", file_path: str = None) -> Dict:
        """Get explanation for a method/class from loaded metadata."""
        key = f"{class_name}.{method_name}"
        
        if key in self.class_explanations:
            return self.class_explanations[key]
        
        # Try alternative key formats
        alternative_keys = [
            class_name,
            method_name,
            f"{class_name}.",
        ]
        
        if file_path:
            module_name = Path(file_path).stem
            alternative_keys.extend([
                f"{module_name}.{class_name}",
                f"{module_name}.{method_name}",
            ])
        
        for alt_key in alternative_keys:
            if alt_key and alt_key in self.class_explanations:
                return self.class_explanations[alt_key]
        
        self.missing_metadata.append({
            "identifier": key,
            "type": item_type,
            "class_name": class_name,
            "item_name": method_name,
            "file": file_path or ""
        })
        return None

    def _infer_role_in_system(self, class_name: str, method_name: str, filename: str) -> str:
        """Infer the role of a class/method within the overall system architecture."""
        # Map classes to their system roles
        class_roles = {
            'KaplanMeierEstimator': 'ML analytics layer — Models lead conversion timing using survival analysis (Kaplan-Meier formula), helping users understand how long leads typically take to convert',
            'LinearRegressionModel': 'ML analytics layer — Predicts lead conversion likelihood using least squares regression with matrix algebra, providing R² accuracy metrics',
            'NaiveBayesCallPredictor': 'ML analytics layer — Predicts optimal call times using Bayesian probability with Laplace smoothing, analysing historical answer patterns by day and hour',
            'WilsonScoreCalculator': 'Statistical analysis layer — Computes confidence intervals for call success rates using the Wilson Score formula, accounting for sample size uncertainty',
            'NLPService': 'NLP processing layer — Provides natural language processing pipeline (tokenisation, stemming, TF-IDF, n-grams) for summarising contact note histories',
            'ContactPriorityQueue': 'Data structure layer — Min-heap priority queue that automatically ranks contacts by calling urgency using a weighted scoring formula',
            'ContactPriorityCalculator': 'Data structure layer — Calculates contact priority scores from Firestore call history data and builds the priority queue',
            'PrioritizedContact': 'Data structure layer — Represents a contact with its calculated priority score for queue ordering',
            'Trie': 'Data structure layer — Prefix tree enabling O(k) autocomplete search across contact notes for fast keyword lookup',
            'TrieNode': 'Data structure layer — Individual node in the Trie, storing children, document IDs, and frequency counts',
            'CallCycleService': 'Call management layer — Tracks the lifecycle of call cycles including starting, pausing, resuming, and reconciling calls with native call history',
            'ExcelImportService': 'Data import layer — Handles Excel/CSV file parsing, duplicate detection, and contact upload to Firebase',
            'TeamService': 'Team management layer — Manages team creation, join codes, member lists, and team-level data access for enterprise use',
            'AuthBloc': 'Authentication layer — BLoC pattern state management for authentication flow, dispatching events and emitting states',
            'FirebaseAuthProvider': 'Authentication layer — Concrete Firebase Auth implementation handling sign-up, sign-in, verification, and error mapping',
            'AuthService': 'Authentication layer — Facade/wrapper around AuthProvider for simplified auth operations',
            'AuthUser': 'Authentication layer — Immutable user model wrapping Firebase User with verified email status',
            'FirebaseCloudStorage': 'Cloud storage layer — Singleton Firestore service for CRUD operations on the notes collection',
            'AppUser': 'Data model — Represents user profiles stored in Firestore with team membership and role information',
            'Team': 'Data model — Represents team entities with join codes and member lists in Firestore',
            'CloudNote': 'Data model — Immutable representation of a note document in Firestore',
            'CloudList': 'Data model — Immutable representation of a contact list entry in Firestore',
        }
        
        if class_name in class_roles:
            base_role = class_roles[class_name]
        elif 'View' in class_name or 'Screen' in class_name or 'Page' in class_name:
            base_role = f'UI layer — Provides the user interface for {class_name.replace("View", "").replace("Screen", "").replace("Page", "")} functionality'
        elif 'State' in class_name:
            base_role = f'State management — Manages the mutable state for the {class_name.replace("State", "").replace("_", "")} widget'
        elif 'Dialog' in class_name:
            base_role = f'UI layer — Modal dialog providing user interaction for {class_name.replace("Dialog", "")} operations'
        elif 'Service' in class_name:
            base_role = f'Service layer — Provides backend operations for {class_name.replace("Service", "")} functionality'
        elif 'Provider' in class_name:
            base_role = f'Provider layer — Supplies {class_name.replace("Provider", "")} data/operations to the widget tree'
        else:
            base_role = f'Utility — Provides helper functionality for the {class_name} module'
        
        return base_role

    def _generate_how_applied(self, purpose: str, class_name: str, method_name: str) -> str:
        """Generate a 'How Applied' description based on context."""
        if not purpose:
            return f"This method in {class_name} implements the functionality described by its code comments."
        
        purpose_lower = purpose.lower()
        
        if 'heap' in purpose_lower or 'priority queue' in purpose_lower:
            return "A min-heap priority queue is implemented to dynamically rank contacts by calling urgency, providing O(log n) insertion and extraction — this is a complex user-defined data structure."
        elif 'trie' in purpose_lower or 'prefix' in purpose_lower:
            return "A Trie (prefix tree) provides O(k) lookup time for autocomplete search across contact notes, significantly faster than linear scanning of all words."
        elif 'kaplan' in purpose_lower or 'survival' in purpose_lower:
            return "Kaplan-Meier survival analysis models lead conversion timing with censored data handling — an advanced statistical technique applied to CRM analytics."
        elif 'regression' in purpose_lower or 'least squares' in purpose_lower or 'matrix' in purpose_lower:
            return "Linear regression using matrix algebra (least squares method) predicts lead conversion from multiple features, with R² calculation to validate model accuracy."
        elif 'naive bayes' in purpose_lower or 'bayes' in purpose_lower or 'laplace' in purpose_lower:
            return "Naive Bayes classification with Laplace smoothing predicts optimal call times by analysing conditional probabilities across day-of-week and hour-of-day combinations."
        elif 'wilson' in purpose_lower or 'confidence' in purpose_lower:
            return "Wilson Score intervals provide statistically robust success rate estimates that account for sample size, preventing misleading metrics from contacts with few calls."
        elif 'nlp' in purpose_lower or 'tokeniz' in purpose_lower or 'tf-idf' in purpose_lower or 'stem' in purpose_lower:
            return "NLP pipeline techniques (tokenisation, stop-word removal, stemming, TF-IDF) are applied to generate automated summaries of contact call note histories."
        elif 'reconcil' in purpose_lower or 'scoring' in purpose_lower:
            return "A weighted multi-factor scoring algorithm matches native call history entries to cycle events, compensating for iOS limitations on direct call state access."
        elif 'firebase' in purpose_lower or 'firestore' in purpose_lower:
            return "Cloud-based data persistence via Firebase Firestore enables real-time synchronisation across devices, addressing the cross-device sync requirement."
        elif 'bloc' in purpose_lower or 'state management' in purpose_lower:
            return "BLoC pattern separates business logic from UI, ensuring testable and maintainable state management for the authentication flow."
        elif 'import' in purpose_lower and ('excel' in purpose_lower or 'csv' in purpose_lower):
            return "File parsing with automatic column detection and duplicate checking enables importing contacts from Excel/CSV, a key user requirement from the analysis phase."
        elif 'team' in purpose_lower or 'join code' in purpose_lower:
            return "Team management with unique join codes supports multi-user enterprise workflows, directly addressing the client requirement for team-level visibility."
        else:
            return f"This implementation in {class_name}.{method_name}() demonstrates the technique: {purpose[:200]}"

    # =========================================================================
    # DOCUMENT GENERATION — TITLE & FRONT MATTER
    # =========================================================================

    def add_title_page(self):
        """Create the title page for the documentation."""
        self.doc.add_paragraph('\n' * 3)
        
        title = self.doc.add_paragraph('NEA Code Documentation')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(28)
            run.font.bold = True
        
        self.doc.add_paragraph()
        
        subtitle = self.doc.add_paragraph('Auto-Dialer Mobile Application')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.size = Pt(18)
        
        self.doc.add_paragraph('\n' * 2)
        
        info = self.doc.add_paragraph('AQA A-Level Computer Science')
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info2 = self.doc.add_paragraph('Non-Exam Assessment (NEA)')
        info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        tech = self.doc.add_paragraph('Built with Flutter/Dart • Firebase Backend')
        tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in tech.runs:
            run.font.size = Pt(12)
            run.italic = True
        
        self.doc.add_paragraph()
        
        date_para = self.doc.add_paragraph(f'Generated: {datetime.now().strftime("%d %B %Y")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph('\n' * 3)
        
        overview = self._add_paragraph(
            "This document provides comprehensive technical documentation of the Auto-Dialer "
            "mobile application source code, including detailed explanations of algorithms, "
            "data structures, machine learning models, and implementation techniques mapped "
            "to AQA A-Level Computer Science assessment criteria.",
            italic=True
        )
        overview.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_page_break()

    def add_table_of_contents(self):
        """Add a table of contents placeholder."""
        self._add_heading('Table of Contents', 2)
        
        toc_para = self.doc.add_paragraph()
        toc_para.add_run('(Table of Contents — Update field after document generation)')
        toc_para.runs[0].italic = True
        
        self.doc.add_page_break()

    def add_aqa_skills_mapping_section(self):
        """Add the AQA Technical Skills Mapping reference section."""
        self._add_heading('AQA Technical Skills Reference', 2)
        
        self._add_paragraph(
            "The following table summarises the AQA A-Level Computer Science technical skills "
            "demonstrated throughout this codebase. Each class and method in the documentation "
            "is mapped to specific skills from these groups."
        )
        
        self.doc.add_paragraph()
        
        skills_data = [
            ("Group A — Complex Skills",
             "• Complex user-defined use of OOP (classes, inheritance, composition, polymorphism, interfaces)\n"
             "• Complex user-defined algorithms (optimisation, minimisation, scheduling, pattern matching)\n"
             "• Recursive algorithms\n"
             "• Graph/Tree traversal\n"
             "• Mergesort or similarly efficient sort\n"
             "• Stack/Queue operations\n"
             "• Hashing\n"
             "• Linked list maintenance\n"
             "• Dynamic generation of objects based on complex OOP model\n"
             "• Calling parameterised Web service APIs and parsing JSON\n"
             "• Advanced matrix operations\n"
             "• Complex client-server model"),
            ("Group B — Intermediate Skills",
             "• Simple user-defined algorithms (mathematical/statistical calculations)\n"
             "• Simple OOP model\n"
             "• Multi-dimensional arrays, Dictionaries, Records\n"
             "• Binary search, Bubble sort\n"
             "• Writing and reading from files\n"
             "• Generation of objects based on simple OOP model\n"
             "• Server-side scripting for client-server model\n"
             "• Calling Web service APIs and parsing JSON"),
            ("Group C — Foundation Skills",
             "• Appropriate choice of simple data types\n"
             "• Single-dimensional arrays\n"
             "• Linear search\n"
             "• Simple mathematical calculations (e.g., average)")
        ]
        
        table = self.doc.add_table(rows=len(skills_data), cols=2)
        table.style = 'Table Grid'
        
        for i, (group, skills) in enumerate(skills_data):
            group_cell = table.rows[i].cells[0]
            skills_cell = table.rows[i].cells[1]
            
            if 'A' in group.split('—')[0]:
                self._set_cell_background(group_cell, "C6EFCE")
            elif 'B' in group.split('—')[0]:
                self._set_cell_background(group_cell, "FFEB9C")
            else:
                self._set_cell_background(group_cell, "BDD7EE")
            
            self._set_cell_width(group_cell, 5.0)
            group_para = group_cell.paragraphs[0]
            group_run = group_para.add_run(group)
            group_run.bold = True
            
            self._set_cell_width(skills_cell, 11.0)
            skills_cell.paragraphs[0].add_run(skills)
        
        self.doc.add_page_break()

    def add_skills_evidence_index(self):
        """Add a skills evidence index showing where each skill is demonstrated."""
        self._add_heading('Skills Evidence Index', 2)
        
        self._add_paragraph(
            "The following index shows where each AQA markscheme skill is demonstrated in the codebase. "
            "Look up the method name in the detailed documentation sections below.",
            italic=True
        )
        
        self.doc.add_paragraph()
        
        groups = {'A': {}, 'B': {}, 'C': {}}
        
        for key, data in self.class_explanations.items():
            group = data.get('group', '').upper()
            skill = data.get('skill_wording', '')
            
            if not skill or skill == 'N/A':
                continue
            
            for g in ['A', 'B', 'C']:
                if g in group:
                    if skill not in groups[g]:
                        groups[g][skill] = []
                    groups[g][skill].append(key)
                    break
        
        group_names = {'A': 'Group A Skills (Complex)', 'B': 'Group B Skills (Intermediate)', 'C': 'Group C Skills (Foundation)'}
        group_colours = {'A': "C6EFCE", 'B': "FFEB9C", 'C': "BDD7EE"}
        
        for g in ['A', 'B', 'C']:
            if groups[g]:
                self._add_heading(group_names[g], 3)
                for skill, methods in sorted(groups[g].items()):
                    skill_table = self.doc.add_table(rows=1, cols=2)
                    skill_table.style = 'Table Grid'
                    
                    skill_cell = skill_table.rows[0].cells[0]
                    methods_cell = skill_table.rows[0].cells[1]
                    
                    self._set_cell_background(skill_cell, group_colours[g])
                    self._set_cell_width(skill_cell, 6.0)
                    skill_para = skill_cell.paragraphs[0]
                    skill_run = skill_para.add_run(skill)
                    skill_run.bold = True
                    
                    self._set_cell_width(methods_cell, 10.0)
                    methods_text = ', '.join(methods[:10])
                    if len(methods) > 10:
                        methods_text += f" (+{len(methods) - 10} more)"
                    methods_cell.paragraphs[0].add_run(methods_text)
                    
                    self.doc.add_paragraph()
        
        self.doc.add_page_break()

    def add_project_objectives_section(self):
        """Add a project objectives summary section."""
        self._add_heading('Project Objectives Summary', 2)
        
        self._add_paragraph(
            "The following table summarises the project objectives and where they are implemented. "
            "For detailed implementation evidence, see the OBJECTIVE field in each method's documentation table.",
            italic=True
        )
        
        self.doc.add_paragraph()
        
        # Key objectives mapped to implementations
        objectives = [
            ("1.1", "User registration/login with Firebase Auth", "AuthBloc, FirebaseAuthProvider, LoginScreen1, RegisterScreen1"),
            ("1.2", "Cloud data sync via Firestore (filtered by user ID)", "FirebaseCloudStorage, firebase_services.dart, all Firestore CRUD operations"),
            ("1.3", "Team/enterprise login with join codes", "TeamService.createTeam, TeamManagementView, AppUser.role"),
            ("1.4", "Team owner dashboard and member management", "TeamManagementView, TeamReportsView, TeamService"),
            ("2.1", "Create/rename/delete contact lists synced to Firestore", "list_view_visible.dart, firebase_services.dart"),
            ("2.2", "Import contacts from Excel/CSV, phone book, manual entry", "ExcelImportService, ContactDirectoryView, contact_upload.dart"),
            ("2.3", "Duplicate detection using normalised phone numbers", "ExcelImportService.checkForDuplicates, normalizePhone()"),
            ("2.4", "Drag-and-drop contact reordering persisted to Firestore", "list_view_visible.dart ReorderableListView, contact_index field"),
            ("3.1", "Auto-dialer via tel:// URI scheme", "dialer.dart, url_launcher integration"),
            ("3.2", "Call cycle tracking with resume capability", "CallCycleService, list_cycles collection, dialer progress persistence"),
            ("3.3", "Call cycle event logging (cycle_events subcollection)", "CallCycleService.recordCallDialogShown/recordDialPressed"),
            ("3.4", "Reconciliation algorithm with weighted scoring", "CallCycleService.reconcileCycle (65% phone, 20% time, 10% order, 5% direction)"),
            ("3.5", "WidgetsBindingObserver for dialer return detection", "dialer.dart didChangeAppLifecycleState"),
            ("4.1", "Post-call feedback dialog with star rating and notes", "CallFeedbackDialog, dialer.dart feedback flow"),
            ("4.2", "Contact notes stored per-contact in Firestore", "contact_notes collection, ContactNotesView"),
            ("4.4", "NLP-generated summary of previous notes before calls", "NLPService.generateNoteSummary, CallFeedbackDialog"),
            ("5.1", "Multi-tab reports dashboard (6 categories)", "ReportsView (6 TabController tabs)"),
            ("5.2", "Activity heatmap with configurable ranges", "reports_view_clean.dart heatmap tab"),
            ("5.3", "Weekly trends chart (line/bar)", "WeeklyCallsChart"),
            ("5.4", "Call duration percentile chart (p25/p50/p75)", "CallDurationChart"),
            ("5.5", "Call outcome donut chart", "CallOutcomeDonutChart"),
            ("5.6", "List performance comparison chart", "ListPerformanceChart"),
            ("5.7", "Team reports viewable per member", "TeamReportsView with targetUserId"),
            ("6.1", "Naive Bayes classifier for optimal call times", "NaiveBayesCallPredictor, CallPredictionView"),
            ("6.2", "Linear regression with R² (least squares, matrix algebra)", "LinearRegressionModel, LinearRegressionStatsView"),
            ("6.3", "Wilson Score confidence intervals", "WilsonScoreCalculator, LinearRegressionStatsView"),
            ("6.4", "Kaplan-Meier survival analysis", "KaplanMeierEstimator, LinearRegressionStatsView"),
            ("6.5", "NLP pipeline (tokenise, stem, TF-IDF, n-grams, summary)", "NLPService (all static methods)"),
            ("6.6", "Contact priority queue (min-heap)", "ContactPriorityQueue, ContactPriorityCalculator"),
            ("6.7", "Trie for autocomplete search", "Trie, ContactNotesView._buildTrieFromNotes"),
        ]
        
        table = self.doc.add_table(rows=len(objectives) + 1, cols=3)
        table.style = 'Table Grid'
        
        headers = table.rows[0].cells
        headers[0].text = "Objective"
        headers[1].text = "Description"
        headers[2].text = "Key Implementation"
        for cell in headers:
            self._set_cell_background(cell, "D9E2F3")
            cell.paragraphs[0].runs[0].bold = True
        
        for i, (obj_id, desc, impl) in enumerate(objectives):
            row = table.rows[i + 1]
            row.cells[0].text = obj_id
            row.cells[1].text = desc
            row.cells[2].text = impl
        
        self.doc.add_paragraph()
        self._add_paragraph(
            "TIP: Use Ctrl+F in Word to search for 'Objective' to find all objective evidence in method tables.",
            italic=True
        )
        
        self.doc.add_page_break()

    # =========================================================================
    # DOCUMENT GENERATION — FILE DOCUMENTATION
    # =========================================================================

    def document_dart_file(self, filepath: Path, doc_style: str = "full"):
        """
        Document a Dart source file.
        
        Args:
            filepath: Path to the Dart file.
            doc_style: "full" for NEA tables, "view" for simplified, "brief" for minimal.
        """
        filename = filepath.name
        relative_path = filepath.relative_to(self.project_root)
        
        # Skip commented-out files
        if filename in self.COMMENTED_OUT_FILES:
            self._add_heading(f'File: {filename} (Legacy — Commented Out)', 3)
            self._add_paragraph(
                f"This file ({relative_path}) contains entirely commented-out legacy code that has been "
                "replaced by the Firebase Firestore cloud-based implementation. It is retained in the "
                "project for reference but is no longer active.",
                italic=True
            )
            self.doc.add_paragraph()
            return
        
        # Read source code
        with open(filepath, 'r', encoding='utf-8') as f:
            source_code = f.read()
            source_lines = source_code.split('\n')
        
        style_label = {
            "full": "Full NEA Tables with AQA Skill Mapping",
            "view": "Widget/View Documentation",
            "brief": "Brief File-Level Table"
        }
        
        self._add_heading(f'Module: {filename}', 3)
        
        # Module overview table
        overview_rows = [
            ("File Path", str(relative_path)),
            ("File Type", self._classify_file_type(filepath)),
            ("Documentation Style", style_label.get(doc_style, "Standard")),
            ("Lines of Code", str(len(source_lines))),
        ]
        
        # Add file-level metadata if available
        file_key = filename.lower()
        if file_key in self.file_metadata:
            fm = self.file_metadata[file_key]
            if fm.get('file_purpose'):
                overview_rows.insert(2, ("File Purpose", fm['file_purpose']))
            if fm.get('role_in_project'):
                overview_rows.append(("Role in Project", fm['role_in_project']))
        
        self._create_vertical_table(overview_rows)
        self.doc.add_paragraph()
        
        # Document imports
        imports = DartParser.extract_imports(source_code)
        if imports:
            self._document_imports(imports, filename)
        
        # Parse and document classes
        classes = DartParser.extract_classes(source_code)
        class_ranges = [(c['lineno'], c['end_lineno']) for c in classes]
        
        # Parse enums and extensions
        enums = DartParser.extract_enums(source_code)
        extensions = DartParser.extract_extensions(source_code)
        
        # Parse top-level functions
        top_funcs = DartParser.extract_top_level_functions(source_code, class_ranges)
        
        # Document enums
        for enum_info in enums:
            self._document_enum(enum_info, filename)
        
        # Document extensions
        for ext_info in extensions:
            self._document_extension(ext_info, filename)
        
        # Document classes
        for class_info in classes:
            if doc_style == "full":
                self._document_class_full(class_info, source_lines, filename)
            elif doc_style == "view":
                self._document_class_view(class_info, source_lines, filename)
            else:
                self._document_class_brief(class_info, filename)
        
        # Document top-level functions
        if top_funcs:
            self._add_heading('Top-Level Functions', 4)
            for func_info in top_funcs:
                self._document_function(func_info, source_lines, filename, doc_style)
        
        self.doc.add_page_break()

    def _classify_file_type(self, filepath: Path) -> str:
        """Classify a Dart file by its type/role."""
        filename = filepath.name
        rel = str(filepath.relative_to(self.project_root))
        
        if filename in self.ML_FILES:
            return "Machine Learning / Statistical Analysis Module"
        elif filename in self.DATA_STRUCTURE_FILES:
            return "Data Structure / Algorithm Module"
        elif filename in self.SERVICE_FILES:
            return "Backend Service Module"
        elif filename in self.MODEL_FILES:
            return "Data Model"
        elif any(p in rel for p in self.VIEW_FILES_PATTERNS):
            return "UI Widget / View"
        elif any(p in rel for p in self.UTILITY_FILES_PATTERNS):
            return "Utility / Helper Module"
        elif 'bloc' in rel:
            return "State Management (BLoC)"
        else:
            return "Dart Module"

    def _document_imports(self, imports: List[Dict], filename: str):
        """Document import statements with explanations."""
        self._add_heading('Import Statements', 4)
        self._add_paragraph(
            "The following imports bring in required packages, libraries, and local modules. "
            "Each import serves a specific purpose in enabling this file's functionality.",
            italic=True
        )
        
        import_code = '\n'.join(imp['line'] for imp in imports)
        
        import_table = self._create_vertical_table([
            ("Section", "Module Imports"),
            ("Line Range", f"Lines {imports[0]['lineno']}-{imports[-1]['lineno']}"),
            ("Import Count", str(len(imports))),
            ("Purpose", "Dependency management — brings in required functionality from Flutter SDK, Firebase, and local modules"),
        ])
        
        self.doc.add_paragraph()
        
        # Check for import explanations in file metadata
        file_key = filename.lower()
        if file_key in self.file_metadata and self.file_metadata[file_key].get('imports_explained'):
            imports_explained = self.file_metadata[file_key]['imports_explained']
            self._add_paragraph("Import Explanations:", bold=True)
            
            explain_table = self.doc.add_table(rows=len(imports_explained) + 1, cols=2)
            explain_table.style = 'Table Grid'
            
            header_cells = explain_table.rows[0].cells
            self._set_cell_background(header_cells[0], "D9E2F3")
            self._set_cell_background(header_cells[1], "D9E2F3")
            header_cells[0].paragraphs[0].add_run("Import").bold = True
            header_cells[1].paragraphs[0].add_run("Purpose / Explanation").bold = True
            
            for i, (import_name, explanation) in enumerate(imports_explained.items(), 1):
                row_cells = explain_table.rows[i].cells
                row_cells[0].paragraphs[0].add_run(import_name)
                row_cells[1].paragraphs[0].add_run(explanation)
            
            self.doc.add_paragraph()
        
        self._add_paragraph("Import Code:", bold=True)
        code_table = self.doc.add_table(rows=1, cols=1)
        code_table.style = 'Table Grid'
        code_cell = code_table.rows[0].cells[0]
        self._set_cell_background(code_cell, "F8F8F8")
        self._add_code_to_cell(code_cell, import_code)
        self.doc.add_paragraph()

    def _document_enum(self, enum_info: Dict, filename: str):
        """Document an enum definition."""
        self._add_heading(f'Enum: {enum_info["name"]}', 4)
        
        rows = [
            ("Enum Name", enum_info['name']),
            ("Source File", filename),
            ("Line Range", f"Lines {enum_info['lineno']}-{enum_info['end_lineno']}"),
            ("Description", enum_info.get('docstring') or 'Enumeration type — see values in code below'),
        ]
        
        self._create_vertical_table(rows)
        self.doc.add_paragraph()
        
        self._add_paragraph("Enum Definition:", bold=True)
        code_table = self.doc.add_table(rows=1, cols=1)
        code_table.style = 'Table Grid'
        code_cell = code_table.rows[0].cells[0]
        self._set_cell_background(code_cell, "F8F8F8")
        self._add_code_to_cell(code_cell, enum_info['source'])
        self.doc.add_paragraph()

    def _document_extension(self, ext_info: Dict, filename: str):
        """Document a Dart extension."""
        self._add_heading(f'Extension: {ext_info["name"]} on {ext_info["on_type"]}', 4)
        
        rows = [
            ("Extension Name", ext_info['name']),
            ("Extends Type", ext_info['on_type']),
            ("Source File", filename),
            ("Description", ext_info.get('docstring') or 'Extension methods — see code below'),
        ]
        
        self._create_vertical_table(rows)
        self.doc.add_paragraph()
        
        self._add_paragraph("Extension Code:", bold=True)
        code_table = self.doc.add_table(rows=1, cols=1)
        code_table.style = 'Table Grid'
        code_cell = code_table.rows[0].cells[0]
        self._set_cell_background(code_cell, "F8F8F8")
        self._add_code_to_cell(code_cell, ext_info['source'])
        self.doc.add_paragraph()

    def _document_class_full(self, class_info: Dict, source_lines: List[str], filename: str):
        """Document a class with full NEA-style tables (for services, ML, data structures)."""
        class_name = class_info['name']
        
        self._add_heading(f'Class: {class_name}', 4)
        
        # Class overview
        class_role = self._infer_role_in_system(class_name, '', filename)
        
        overview_rows = [
            ("Class Name", class_name),
            ("Source File", filename),
            ("Abstract", "Yes" if class_info.get('is_abstract') else "No"),
            ("Role in System", class_role),
            ("Methods Count", str(len(class_info.get('methods', [])))),
        ]
        
        if class_info.get('docstring'):
            overview_rows.insert(3, ("Description", class_info['docstring']))
        if class_info.get('extends'):
            overview_rows.insert(2, ("Extends", class_info['extends']))
        if class_info.get('with_mixins'):
            overview_rows.insert(3, ("Mixins", class_info['with_mixins']))
        if class_info.get('implements'):
            overview_rows.insert(3, ("Implements", class_info['implements']))
        if class_info.get('fields'):
            overview_rows.append(("Class Fields", ', '.join(class_info['fields'][:15])))
        
        self._create_vertical_table(overview_rows)
        self.doc.add_paragraph()
        
        # Document each method
        for method_info in class_info.get('methods', []):
            self._document_method_full(class_name, method_info, source_lines, filename)

    def _document_method_full(self, class_name: str, method_info: Dict,
                               source_lines: List[str], filename: str):
        """Document a method with full NEA table format."""
        method_name = method_info['name']
        
        self._add_heading(f'Method: {method_name}', 5)
        
        # Get explanation from metadata
        # For constructors, try ClassName.ClassName or just ClassName
        lookup_name = method_name
        if method_info.get('type') in ('constructor', 'factory_constructor'):
            lookup_name = method_name.split('.')[-1] if '.' in method_name else method_name
        
        explanation = self._get_method_explanation(class_name, lookup_name, method_info.get('type', 'method'), filename)
        
        # Prepare parameters string
        params_str = method_info.get('params_raw', 'None') or 'None'
        if explanation and explanation.get('parameters_detailed'):
            param_lines = []
            for param in explanation['parameters_detailed']:
                ptype = f" ({param.get('type', '')})" if param.get('type') and param.get('type') != 'dynamic' else ""
                exp_text = param.get('explanation', '')
                param_lines.append(f"• {param['name']}{ptype}: {exp_text}")
            if param_lines:
                params_str = '\n'.join(param_lines)
        
        # Get group colour for highlighting
        group = explanation.get('group', '') if explanation else ''
        group_colour = self._get_group_colour(group)
        
        highlight_rows = {}
        if group:
            highlight_rows["Markscheme Group"] = group_colour
            highlight_rows["Markscheme Skill"] = group_colour
            highlight_rows["How Applied"] = group_colour
        
        role_in_system = self._infer_role_in_system(class_name, method_name, filename)
        
        if explanation:
            purpose_text = explanation.get('purpose', '')
            if not purpose_text:
                purpose_text = method_info.get('docstring') or 'See code implementation and inline comments'
            
            rows = [
                ("Identifier", f"{class_name}.{method_name}"),
                ("Purpose", purpose_text),
                ("Role in System", role_in_system),
                ("Parameters", params_str),
                ("Returns", "See implementation"),
            ]
            
            if explanation.get('objective'):
                rows.append(("Project Objective", explanation['objective']))
            
            group_text = explanation.get('group', 'N/A')
            skill_text = explanation.get('skill_wording', 'N/A')
            how_applied = explanation.get('how_applied', '')
            
            if not how_applied or how_applied == 'See code implementation':
                how_applied = self._generate_how_applied(purpose_text, class_name, method_name)
            
            rows.extend([
                ("Markscheme Group", group_text if group_text else "Not yet classified"),
                ("Markscheme Skill", skill_text if skill_text else "Not yet classified"),
                ("How Applied", how_applied),
            ])
        else:
            purpose_text = method_info.get('docstring') or 'See code implementation'
            
            rows = [
                ("Identifier", f"{class_name}.{method_name}"),
                ("Purpose", f"[METADATA NEEDED] {purpose_text}"),
                ("Role in System", role_in_system),
                ("Parameters", params_str),
                ("Returns", "See implementation"),
                ("Markscheme Group", "[Add to class_explanations_extended.txt]"),
                ("Markscheme Skill", "[Add to class_explanations_extended.txt]"),
                ("How Applied", "[Add to class_explanations_extended.txt]"),
            ]
        
        table = self._create_vertical_table(rows, highlight_rows=highlight_rows)
        self.doc.add_paragraph()
        
        # Add the actual code
        self._add_paragraph("Method Implementation:", bold=True)
        code_table = self.doc.add_table(rows=1, cols=1)
        code_table.style = 'Table Grid'
        code_cell = code_table.rows[0].cells[0]
        self._set_cell_background(code_cell, "F8F8F8")
        self._add_code_to_cell(code_cell, method_info.get('source', ''))
        
        self.doc.add_paragraph()

    def _document_class_view(self, class_info: Dict, source_lines: List[str], filename: str):
        """Document a view/widget class with simplified tables."""
        class_name = class_info['name']
        
        self._add_heading(f'Widget: {class_name}', 4)
        
        # Get explanation
        explanation = self._get_method_explanation(class_name, '', 'widget', filename)
        
        purpose = ''
        if explanation:
            purpose = explanation.get('purpose', '')
        if not purpose:
            purpose = class_info.get('docstring') or f'UI widget providing {class_name} functionality'
        
        overview_rows = [
            ("Widget Name", class_name),
            ("Source File", filename),
            ("Purpose", purpose),
            ("Type", "StatefulWidget" if class_info.get('extends') and 'Stateful' in (class_info['extends'] or '') else "StatelessWidget"),
        ]
        
        if class_info.get('extends'):
            overview_rows.append(("Extends", class_info['extends']))
        if class_info.get('with_mixins'):
            overview_rows.append(("Mixins", class_info['with_mixins']))
        
        method_names = [m['name'] for m in class_info.get('methods', []) if not m['name'].startswith('_')]
        if method_names:
            overview_rows.append(("Public Methods", ', '.join(method_names[:10])))
        
        # Add objective if present
        if explanation and explanation.get('objective'):
            overview_rows.append(("Project Objective", explanation['objective']))
        
        # Add markscheme if present
        if explanation and explanation.get('group'):
            group_colour = self._get_group_colour(explanation['group'])
            overview_rows.extend([
                ("Markscheme Group", explanation.get('group', 'N/A')),
                ("Markscheme Skill", explanation.get('skill_wording', 'N/A')),
                ("How Applied", explanation.get('how_applied', 'See implementation')),
            ])
        
        self._create_vertical_table(overview_rows)
        self.doc.add_paragraph()
        
        # Add the class code (full)
        self._add_paragraph("Widget Implementation:", bold=True)
        code_table = self.doc.add_table(rows=1, cols=1)
        code_table.style = 'Table Grid'
        code_cell = code_table.rows[0].cells[0]
        self._set_cell_background(code_cell, "F8F8F8")
        self._add_code_to_cell(code_cell, class_info.get('source', ''))
        
        self.doc.add_paragraph()

    def _document_class_brief(self, class_info: Dict, filename: str):
        """Document a class with brief file-level table."""
        class_name = class_info['name']
        
        self._add_heading(f'Class: {class_name}', 4)
        
        rows = [
            ("Class Name", class_name),
            ("Source File", filename),
            ("Description", class_info.get('docstring') or 'See code implementation'),
        ]
        
        if class_info.get('fields'):
            rows.append(("Fields", ', '.join(class_info['fields'][:10])))
        
        method_names = [m['name'] for m in class_info.get('methods', [])]
        if method_names:
            rows.append(("Methods", ', '.join(method_names[:10])))
        
        self._create_vertical_table(rows)
        self.doc.add_paragraph()
        
        # Add code
        self._add_paragraph("Class Code:", bold=True)
        code_table = self.doc.add_table(rows=1, cols=1)
        code_table.style = 'Table Grid'
        code_cell = code_table.rows[0].cells[0]
        self._set_cell_background(code_cell, "F8F8F8")
        self._add_code_to_cell(code_cell, class_info.get('source', ''))
        
        self.doc.add_paragraph()

    def _document_function(self, func_info: Dict, source_lines: List[str], filename: str, doc_style: str):
        """Document a top-level function."""
        func_name = func_info['name']
        
        self._add_heading(f'Function: {func_name}', 5)
        
        explanation = self._get_method_explanation(Path(filename).stem, func_name, "function", filename)
        
        rows = [
            ("Identifier", func_name),
            ("Purpose", explanation.get('purpose', func_info.get('docstring', 'See code comments')) if explanation else (func_info.get('docstring') or 'See code comments')),
            ("Parameters", func_info.get('params_raw', 'None') or 'None'),
        ]
        
        if explanation and explanation.get('objective'):
            rows.append(("Project Objective", explanation['objective']))
        if explanation and explanation.get('group'):
            rows.extend([
                ("Markscheme Group", explanation.get('group', 'N/A')),
                ("Markscheme Skill", explanation.get('skill_wording', 'N/A')),
                ("How Applied", explanation.get('how_applied', 'See implementation')),
            ])
        
        self._create_vertical_table(rows)
        self.doc.add_paragraph()
        
        self._add_paragraph("Function Implementation:", bold=True)
        code_table = self.doc.add_table(rows=1, cols=1)
        code_table.style = 'Table Grid'
        code_cell = code_table.rows[0].cells[0]
        self._set_cell_background(code_cell, "F8F8F8")
        self._add_code_to_cell(code_cell, func_info.get('source', ''))
        
        self.doc.add_paragraph()

    # =========================================================================
    # MISSING METADATA REPORT
    # =========================================================================

    def add_missing_metadata_report(self):
        """Add a section reporting any missing metadata."""
        if not self.missing_metadata:
            return
        
        self._add_heading('Missing Metadata Report', 2)
        
        self._add_paragraph(
            "WARNING: The following items were found in the source code but did not have "
            "corresponding entries in the metadata files. These should be added to "
            "lib/ExplanationFiles/class_explanations_extended.txt to ensure complete documentation:",
            bold=True
        )
        
        self.doc.add_paragraph()
        
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = "#"
        header_cells[1].text = "Type"
        header_cells[2].text = "Missing Identifier"
        header_cells[3].text = "File"
        for cell in header_cells:
            self._set_cell_background(cell, "FFD9D9")
        
        seen = set()
        unique_items = []
        for item in self.missing_metadata:
            identifier = item.get("identifier", "")
            if identifier not in seen:
                seen.add(identifier)
                unique_items.append(item)
        
        for i, item in enumerate(sorted(unique_items, key=lambda x: x.get("identifier", "")), 1):
            row = table.add_row()
            row.cells[0].text = str(i)
            row.cells[1].text = item.get("type", "unknown")
            row.cells[2].text = item.get("identifier", "")
            row.cells[3].text = item.get("file", "")
        
        self.doc.add_page_break()

    # =========================================================================
    # MAIN GENERATION PIPELINE
    # =========================================================================

    def _collect_all_dart_files(self) -> List[Path]:
        """Collect all Dart files from the lib/ directory."""
        # If project_root is already 'lib', scan it directly; otherwise look for lib/ inside it
        if self.project_root.name == 'lib':
            lib_dir = self.project_root
        else:
            lib_dir = self.project_root / 'lib'
        files = []
        
        if lib_dir.exists():
            for dart_file in sorted(lib_dir.rglob('*.dart')):
                # Skip auto-generated and build files
                if any(skip in dart_file.name for skip in self.SKIP_FILES):
                    continue
                if 'build/' in str(dart_file) or '.dart_tool' in str(dart_file):
                    continue
                files.append(dart_file)
        
        return files

    def _determine_doc_style(self, filepath: Path) -> str:
        """Determine which documentation style to use for a file."""
        filename = filepath.name
        rel = str(filepath.relative_to(self.project_root))
        
        # Full NEA tables for ML, data structures, and core services
        if filename in self.ML_FILES:
            return "full"
        if filename in self.DATA_STRUCTURE_FILES:
            return "full"
        if filename in self.SERVICE_FILES:
            return "full"
        if 'bloc' in rel:
            return "full"
        
        # View documentation for UI files
        if any(p in rel for p in self.VIEW_FILES_PATTERNS):
            return "view"
        
        # Brief for models, constants, utilities
        if filename in self.MODEL_FILES:
            return "brief"
        if any(p in rel for p in self.UTILITY_FILES_PATTERNS):
            return "brief"
        
        return "full"

    def generate(self, files_to_document: Optional[List[str]] = None):
        """
        Generate the complete documentation.
        
        Args:
            files_to_document: Optional list of specific file paths to document.
        """
        print("Starting NEA Code Documentation Generation (Flutter/Dart)...")
        print("=" * 60)
        
        # Add document structure
        self.add_title_page()
        self.add_table_of_contents()
        self.add_aqa_skills_mapping_section()
        self.add_skills_evidence_index()
        self.add_project_objectives_section()
        
        # Collect files
        if files_to_document:
            files = [self.project_root / f for f in files_to_document]
        else:
            files = self._collect_all_dart_files()
        
        # Categorise files for ordered documentation
        ml_files = []
        ds_files = []
        service_files = []
        bloc_files = []
        model_files = []
        view_files = []
        utility_files = []
        other_files = []
        
        for filepath in files:
            if not filepath.exists():
                print(f"Warning: File not found: {filepath}")
                continue
            
            filename = filepath.name
            rel = str(filepath.relative_to(self.project_root))
            
            if filename in self.ML_FILES:
                ml_files.append(filepath)
            elif filename in self.DATA_STRUCTURE_FILES:
                ds_files.append(filepath)
            elif filename in self.SERVICE_FILES:
                service_files.append(filepath)
            elif 'bloc' in rel:
                bloc_files.append(filepath)
            elif filename in self.MODEL_FILES:
                model_files.append(filepath)
            elif any(p in rel for p in self.VIEW_FILES_PATTERNS):
                view_files.append(filepath)
            elif any(p in rel for p in self.UTILITY_FILES_PATTERNS):
                utility_files.append(filepath)
            else:
                other_files.append(filepath)
        
        # Document each category
        categories = [
            ("MACHINE LEARNING & STATISTICAL ANALYSIS MODULES", ml_files,
             "These modules implement machine learning models, statistical analysis, "
             "and data science algorithms — the most technically complex parts of the codebase."),
            ("DATA STRUCTURES & ALGORITHMS", ds_files,
             "These modules implement custom data structures and algorithms used throughout the application."),
            ("BACKEND SERVICE MODULES", service_files + bloc_files,
             "These modules provide backend services including authentication, cloud storage, "
             "call cycle management, and team management."),
            ("DATA MODELS", model_files,
             "These modules define the data models used throughout the application to represent "
             "entities stored in Firebase Firestore."),
            ("UI VIEWS & WIDGETS", view_files,
             "These modules implement the user interface screens and widgets."),
            ("UTILITIES, HELPERS & THEME", utility_files,
             "These modules provide shared utility functions, dialog helpers, and theme configuration."),
            ("OTHER MODULES", other_files,
             "Additional application modules."),
        ]
        
        for category_name, category_files, description in categories:
            if not category_files:
                continue
            
            self._add_heading(category_name, 1)
            self._add_paragraph(description, italic=True)
            self.doc.add_page_break()
            
            for filepath in sorted(category_files):
                print(f"Documenting: {filepath.relative_to(self.project_root)}")
                doc_style = self._determine_doc_style(filepath)
                try:
                    self.document_dart_file(filepath, doc_style)
                except Exception as e:
                    print(f"  Error documenting {filepath.name}: {e}")
                    self._add_paragraph(f"Error documenting {filepath.name}: {e}", italic=True)
        
        # Add annotations from JSON if available
        self._add_json_annotations()
        
        # Add missing metadata report
        self.add_missing_metadata_report()
        
        # Save document
        output_path = self.project_root / self.output_file
        self.doc.save(output_path)
        
        print("=" * 60)
        print(f"Documentation generated successfully: {output_path}")
        
        if self.missing_metadata:
            seen = set()
            unique_items = []
            for item in self.missing_metadata:
                identifier = item.get("identifier", "")
                if identifier not in seen:
                    seen.add(identifier)
                    unique_items.append(item)
            
            print(f"\nWarning: {len(unique_items)} items missing metadata.")
            print("   See 'Missing Metadata Report' section in the document.")
            
            missing_json_path = self.project_root / "missing_metadata_items.json"
            with open(missing_json_path, 'w', encoding='utf-8') as f:
                json.dump(unique_items, f, indent=2, ensure_ascii=False)
            print(f"   Missing metadata exported to: {missing_json_path}")
        
        return output_path

    def _add_json_annotations(self):
        """Add textbox annotations from JSON metadata if available."""
        annotations_file = self.workspace_root / 'lib' / 'ExplanationFiles' / 'annotations.json'
        
        annotations = []
        if annotations_file.exists():
            with open(annotations_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                annotations = data.get('annotations', []) if isinstance(data, dict) else data
        elif 'annotations' in self.system_context:
            annotations = self.system_context['annotations']
        
        if annotations:
            self._add_heading('Design Annotations', 2)
            self._add_paragraph(
                "The following annotations provide additional context about design decisions, "
                "data flow, and implementation intent for key components of the system.",
                italic=True
            )
            self.doc.add_paragraph()
            
            for annotation in annotations:
                reference = annotation.get('reference', '')
                text = annotation.get('text', '')
                if text:
                    self._add_textbox_annotation(text, reference)


# =============================================================================
# FILE SELECTION & MAIN ENTRY POINT
# =============================================================================

def select_files(all_files: List[Path], project_root: Path) -> List[Path]:
    """Interactive file selection menu."""
    if not all_files:
        print("No files found to document.")
        return []
    
    print("\n" + "=" * 60)
    print("FILE SELECTION")
    print("=" * 60)
    print("\nAvailable Dart files to document:")
    print("-" * 40)
    
    idx = 1
    file_map = {}
    
    for filepath in sorted(all_files):
        rel_path = filepath.relative_to(project_root)
        print(f"  {idx:3}. {rel_path}")
        file_map[idx] = filepath
        idx += 1
    
    print("\n" + "-" * 40)
    print("Selection options:")
    print("  - Enter 'all' or '0' or press Enter to select ALL files")
    print("  - Enter specific numbers separated by commas: 1,3,5")
    print("  - Enter ranges: 1-5,8-10")
    print("  - Enter 'q' to cancel")
    print("-" * 40)
    
    selection = input("\nYour selection: ").strip().lower()
    
    if selection == 'q':
        print("Selection cancelled.")
        return []
    
    if not selection or selection == 'all' or selection == '0':
        print(f"\nSelected ALL {len(all_files)} files.")
        return all_files
    
    selected_indices = set()
    try:
        parts = selection.replace(' ', '').split(',')
        for part in parts:
            if '-' in part:
                start, end = part.split('-', 1)
                selected_indices.update(range(int(start), int(end) + 1))
            else:
                selected_indices.add(int(part))
    except ValueError:
        print("Invalid selection format. Selecting all files.")
        return all_files
    
    selected_files = []
    for idx in sorted(selected_indices):
        if idx in file_map:
            selected_files.append(file_map[idx])
    
    if not selected_files:
        print("No valid files selected. Selecting all files.")
        return all_files
    
    print(f"\nSelected {len(selected_files)} file(s):")
    for f in selected_files:
        print(f"  - {f.relative_to(project_root)}")
    
    return selected_files


def main():
    """Main entry point for the documentation generator."""
    parser = argparse.ArgumentParser(
        description='Generate NEA Code Documentation for AQA A-Level CS (Flutter/Dart)'
    )
    parser.add_argument(
        'directory',
        nargs='?',
        default='.',
        help='Path to project root directory (default: current directory)'
    )
    parser.add_argument(
        '-o', '--output',
        default='NEA_Code_Documentation.docx',
        help='Output filename (default: NEA_Code_Documentation.docx)'
    )
    parser.add_argument(
        '-a', '--append',
        action='store_true',
        help='Append to existing document instead of prompting'
    )
    parser.add_argument(
        '-r', '--replace',
        action='store_true',
        help='Replace existing document (overwrites)'
    )
    parser.add_argument(
        '-s', '--skip-selection',
        action='store_true',
        help='Skip file selection and include all files'
    )
    parser.add_argument(
        '--files',
        nargs='+',
        help='Specific files to document (skips interactive selection)'
    )
    
    args = parser.parse_args()
    
    project_root = Path(args.directory).resolve()
    
    if not project_root.exists():
        print(f"Error: Project root does not exist: {project_root}")
        sys.exit(1)
    
    output_path = project_root / args.output
    
    append_mode = False
    
    if output_path.exists():
        if args.append and args.replace:
            print("Error: Cannot specify both --append and --replace")
            sys.exit(1)
        elif args.append:
            append_mode = True
        elif args.replace:
            append_mode = False
        else:
            print(f"\nOutput file already exists: {output_path}")
            print("\nChoose action:")
            print("  (A)ppend — Add new content to existing document")
            print("  (R)eplace — Create new document (overwrites existing)")
            print("  (C)ancel — Exit without changes")
            
            choice = input("\nYour choice [A/R/C]: ").strip().upper()
            
            if choice == 'C' or not choice:
                print("Operation cancelled.")
                sys.exit(0)
            elif choice == 'A':
                append_mode = True
            elif choice == 'R':
                append_mode = False
            else:
                print("Invalid choice. Operation cancelled.")
                sys.exit(1)
    
    files_to_document = None
    
    if args.files:
        files_to_document = args.files
    elif not args.skip_selection:
        temp_generator = NEACodeDocumentationGenerator(
            project_root=str(project_root),
            output_file=args.output,
            append_mode=append_mode
        )
        all_files = temp_generator._collect_all_dart_files()
        
        selected_files = select_files(all_files, project_root)
        
        if not selected_files:
            print("No files selected. Exiting.")
            sys.exit(0)
        
        files_to_document = [str(f.relative_to(project_root)) for f in selected_files]
    
    generator = NEACodeDocumentationGenerator(
        project_root=str(project_root),
        output_file=args.output,
        append_mode=append_mode
    )
    
    try:
        result_path = generator.generate(files_to_document=files_to_document)
        action = "appended to" if append_mode else "created"
        print(f"\nDocumentation {action}: {result_path}")
    except MissingMetadataError as e:
        print(f"\nError: Missing required metadata")
        print(f"   {e}")
        sys.exit(1)
    except Exception as e:
        print(f"\nError generating documentation: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
